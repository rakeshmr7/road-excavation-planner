import os
import httpx
import uuid
import docx
from pypdf import PdfReader
from chromadb.api.types import EmbeddingFunction, Documents, Embeddings
import chromadb
from langsmith import traceable
from app.core.config import settings
from app.services.failover_llm import query_llm
from app.schemas.ai import RAGQueryResponse, SourceCitation
from typing import List, Dict, Any

# Ensure ChromaDB directory exists
os.makedirs(settings.CHROMA_DB_PATH, exist_ok=True)

def pad_or_truncate(v: List[float], target_len: int = 768) -> List[float]:
    """Ensures vectors have uniform dimension (768) to prevent ChromaDB schema conflicts."""
    if len(v) > target_len:
        return v[:target_len]
    elif len(v) < target_len:
        return v + [0.0] * (target_len - len(v))
    return v

class FailoverEmbeddingFunction(EmbeddingFunction):
    """
    ChromaDB compatible embedding function that queries Gemini API first,
    falls back to local Ollama, and defaults to a mock vector if offline.
    """
    def __call__(self, input: Documents) -> Embeddings:
        embeddings = []
        for text in input:
            embeddings.append(self.get_embedding(text))
        return embeddings

    def get_embedding(self, text: str) -> List[float]:
        # 1. Attempt Gemini Embeddings
        if settings.GEMINI_API_KEY:
            for attempt in range(3):
                try:
                    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-embedding-2:embedContent?key={settings.GEMINI_API_KEY}"
                    payload = {
                        "model": "models/gemini-embedding-2",
                        "content": {"parts": [{"text": text}]}
                    }
                    with httpx.Client() as client:
                        response = client.post(url, json=payload, timeout=15.0)
                        response.raise_for_status()
                        vector = response.json()["embedding"]["values"]
                        return pad_or_truncate(vector)
                except Exception as e:
                    print(f"RAG Embeddings Warning: Gemini embedding attempt {attempt+1} failed: {e}")
                    if attempt < 2:
                        import time
                        time.sleep(0.5 * (attempt + 1))
                    else:
                        print("RAG Embeddings: Gemini embedding failed permanently for this chunk.")

        # 2. Attempt Ollama Embeddings
        try:
            url = f"{settings.OLLAMA_HOST}/api/embeddings"
            payload = {
                "model": "llama3",
                "prompt": text
            }
            with httpx.Client() as client:
                response = client.post(url, json=payload, timeout=30.0)
                response.raise_for_status()
                vector = response.json()["embedding"]
                return pad_or_truncate(vector)
        except Exception as e:
            print(f"RAG Embeddings Warning: Ollama embedding failed: {e}")

        # 3. Safe Mock Fallback (preventing crash during offline development)
        print("RAG Embeddings: Using mock vector fallback for chunk.")
        val = sum(ord(c) for c in text) / 1000.0
        mock_vector = [val * (i % 5 + 1) * 0.01 for i in range(768)]
        return pad_or_truncate(mock_vector)


# Initialize ChromaDB client & collection
chroma_client = chromadb.PersistentClient(path=settings.CHROMA_DB_PATH)
embedding_fn = FailoverEmbeddingFunction()
collection = chroma_client.get_or_create_collection(
    name="gcc_policies",
    embedding_function=embedding_fn
)

def extract_text_from_pdf(file_path: str) -> List[Dict[str, Any]]:
    """Extracts text page by page from a PDF file."""
    pages_data = []
    try:
        with open(file_path, "rb") as f:
            reader = PdfReader(f)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages_data.append({
                        "text": text,
                        "page_number": i + 1
                    })
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return pages_data

def extract_text_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """Extracts text from a DOCX file."""
    doc_data = []
    try:
        doc = docx.Document(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Word documents don't have explicit pages natively, treating full content as page 1
        doc_data.append({
            "text": "\n".join(full_text),
            "page_number": 1
        })
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return doc_data

_paddle_ocr_instance = None

def get_paddle_ocr():
    global _paddle_ocr_instance
    if _paddle_ocr_instance is None:
        from paddleocr import PaddleOCR
        _paddle_ocr_instance = PaddleOCR(lang='en')
    return _paddle_ocr_instance

import re

def clean_ocr_text(lines: List[str]) -> str:
    """
    Cleans OCR lines: removes duplicate spaces, merges wrapped paragraph lines,
    and preserves headings, lists, and numbered clauses.
    """
    cleaned_paragraphs = []
    current_paragraph = []
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # Remove duplicate spaces
        stripped = re.sub(r'\s+', ' ', stripped)
        
        # Identify headings: e.g., "SECTION 1", "POLICY 4.2", "CHAPTER V" or short all-caps lines
        is_heading = (
            re.match(r'^(section|policy|chapter|article|clause|annex)\s+\d+', stripped, re.IGNORECASE) or
            (stripped.isupper() and len(stripped) < 80)
        )
        
        # Identify list item or numbered clause: e.g., "1.", "a)", "•", "-", "1.1.1"
        is_list_or_clause = re.match(r'^(\d+(\.\d+)*\b|[\-\*•\+•]|\([a-zA-Z0-9]+\)|[a-zA-Z]\.)\s+', stripped)
        
        if is_heading or is_list_or_clause:
            # Flush current paragraph first
            if current_paragraph:
                cleaned_paragraphs.append(" ".join(current_paragraph))
                current_paragraph = []
            cleaned_paragraphs.append(stripped)
        else:
            # If it's a regular text line:
            current_paragraph.append(stripped)
    
    if current_paragraph:
        cleaned_paragraphs.append(" ".join(current_paragraph))
        
    return "\n\n".join(cleaned_paragraphs)

def perform_ocr_on_image(img) -> tuple[str, float | None]:
    """
    Runs PaddleOCR on a PIL image. If it fails, falls back to Tesseract.
    Returns (extracted_text, ocr_confidence).
    """
    # 1. Attempt PaddleOCR
    try:
        import numpy as np
        ocr_engine = get_paddle_ocr()
        img_arr = np.array(img)
        # Convert RGB to BGR for OpenCV / PaddleOCR compatibility
        if len(img_arr.shape) == 3 and img_arr.shape[2] == 3:
            img_arr = img_arr[:, :, ::-1]
            
        result = ocr_engine.ocr(img_arr, cls=True)
        
        lines = []
        confidences = []
        if result:
            for line in result:
                if not line:
                    continue
                for word_info in line:
                    text_str, conf = word_info[1]
                    if text_str and text_str.strip():
                        lines.append(text_str)
                        confidences.append(float(conf))
        
        if lines:
            cleaned = clean_ocr_text(lines)
            avg_conf = sum(confidences) / len(confidences) if confidences else None
            return cleaned, avg_conf
    except Exception as e:
        print(f"PaddleOCR failed, falling back to Tesseract: {e}")

    # 2. Fallback to Tesseract OCR
    try:
        import pytesseract
        if not pytesseract.pytesseract.tesseract_cmd or pytesseract.pytesseract.tesseract_cmd == "tesseract":
            for path in [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                r"C:\Users\rakes\AppData\Local\Tesseract-OCR\tesseract.exe",
            ]:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    break
        
        # Pytesseract returns plain text directly
        text = pytesseract.image_to_string(img)
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        cleaned = clean_ocr_text(lines)
        return cleaned, None
    except Exception as e:
        print(f"Tesseract OCR failed as well: {e}")
        raise RuntimeError(f"OCR engines failed: {e}")

def extract_text_via_ocr_pdf(file_path: str) -> list[dict[str, Any]]:
    """
    Renders PDF pages into images and runs OCR on each page.
    Returns a list of dicts with {"text": text, "page_number": page_num, "ocr_confidence": conf}
    """
    import pypdfium2 as pdfium
    pages_data = []
    
    doc = None
    try:
        doc = pdfium.PdfDocument(file_path)
        num_pages = len(doc)
        
        for i in range(num_pages):
            page = doc[i]
            bitmap = page.render(scale=2)
            pil_img = bitmap.to_pil()
            
            try:
                text, conf = perform_ocr_on_image(pil_img)
                if text and text.strip():
                    pages_data.append({
                        "text": text,
                        "page_number": i + 1,
                        "ocr_confidence": conf
                    })
            except Exception as e:
                print(f"Error doing OCR on PDF page {i+1}: {e}")
    finally:
        if doc is not None:
            doc.close()
            
    return pages_data

async def index_policy_document(policy_id: uuid.UUID, file_path: str) -> None:
    """
    Extracts text from PDF/DOCX or Image, splits into chunks, and pushes to ChromaDB collection.
    """
    import time
    file_name = os.path.basename(file_path)
    ext = file_name.split(".")[-1].lower()
    
    extraction_method = "text"
    pages = []
    start_time = time.time()
    
    if ext == "pdf":
        try:
            pages = extract_text_from_pdf(file_path)
            # Check if PDF is scanned (contains no selectable text)
            total_text = "".join(page["text"] for page in pages).strip()
            if not total_text:
                print(f"RAG: PDF '{file_name}' appears to be scanned/image-based. Performing OCR.")
                pages = extract_text_via_ocr_pdf(file_path)
                extraction_method = "ocr"
        except Exception as e:
            print(f"RAG Warning: PDF text extraction failed for '{file_name}'. Retrying with OCR: {e}")
            try:
                pages = extract_text_via_ocr_pdf(file_path)
                extraction_method = "ocr"
            except Exception as ocr_err:
                print(f"RAG Error: OCR extraction fallback failed for '{file_name}': {ocr_err}")
                raise RuntimeError(f"Failed to extract text or perform OCR on PDF '{file_name}': {ocr_err}")
                
    elif ext == "docx":
        pages = extract_text_from_docx(file_path)
        
    elif ext in ["jpg", "jpeg", "png", "tiff"]:
        extraction_method = "ocr"
        try:
            from PIL import Image
            with Image.open(file_path) as img:
                text, conf = perform_ocr_on_image(img)
                if text and text.strip():
                    pages = [{
                        "text": text,
                        "page_number": 1,
                        "ocr_confidence": conf
                    }]
        except Exception as e:
            print(f"RAG Error: Image OCR failed for '{file_name}': {e}")
            raise RuntimeError(f"Failed to perform OCR on image '{file_name}': {e}")
    else:
        raise ValueError("Unsupported extension")

    if not pages:
        raise ValueError(f"No text content could be extracted or read from '{file_name}'.")

    chunks = []
    metadatas = []
    ids = []

    # Chunking: 1000 characters with 200 characters overlap
    chunk_size = 1000
    overlap = 200

    ocr_confidences = []

    for page in pages:
        text = page["text"]
        page_num = page["page_number"]
        ocr_conf = page.get("ocr_confidence")
        if ocr_conf is not None:
            ocr_confidences.append(ocr_conf)
        
        start = 0
        chunk_idx = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            
            chunks.append(chunk)
            metadatas.append({
                "policy_id": str(policy_id),
                "file_name": file_name,
                "page": page_num,
                "chunk_index": chunk_idx,
                "extraction_method": extraction_method,
                "page_number": page_num,
                "ocr_confidence": ocr_conf
            })
            ids.append(f"{policy_id}_{page_num}_{chunk_idx}")
            
            start += (chunk_size - overlap)
            chunk_idx += 1

    if chunks:
        collection.add(
            documents=chunks,
            metadatas=metadatas,
            ids=ids
        )
        processing_time = time.time() - start_time
        num_pages = len(pages)
        avg_ocr_conf = sum(ocr_confidences) / len(ocr_confidences) if ocr_confidences else None
        
        # Logging details
        ocr_conf_log = f" | OCR Confidence: {avg_ocr_conf:.2%}" if avg_ocr_conf is not None else ""
        print(
            f"RAG LOG: [File: '{file_name}'] "
            f"[Extraction Method: {extraction_method}] "
            f"[Pages Processed: {num_pages}] "
            f"[Chunks Created: {len(chunks)}] "
            f"[Processing Time: {processing_time:.2f}s]"
            f"{ocr_conf_log}"
        )
        print(f"RAG: Successfully indexed {len(chunks)} chunks for document '{file_name}' in ChromaDB.")

async def remove_policy_document(policy_id: uuid.UUID) -> None:
    """
    Purges all vector embeddings associated with a policy ID from ChromaDB.
    """
    collection.delete(
        where={"policy_id": str(policy_id)}
    )
    print(f"RAG: Purged vector entries for policy ID: {policy_id}")

from typing import Optional

def format_proposal_for_rag(proposal, analysis, remarks: Optional[str] = None) -> str:
    """
    Formats the proposal details, its AI compliance scan findings,
    and any admin decisions/remarks into a unified text document for ChromaDB indexing.
    """
    remarks_str = f"Admin Decision Remarks: {remarks}" if remarks else "Admin Decision Remarks: None"
    
    text = (
        f"PROPOSAL PLAN DETAILS:\n"
        f"Proposal ID: {proposal.id}\n"
        f"Road Name: {proposal.road_name}\n"
        f"Department: {proposal.department.upper()}\n"
        f"Excavation Purpose: {proposal.purpose}\n"
        f"Technical Description: {proposal.description}\n"
        f"Start Date: {proposal.start_date}\n"
        f"End Date: {proposal.end_date}\n"
        f"Estimated Budget: Rs {proposal.estimated_budget}\n"
        f"Contractor: {proposal.contractor}\n"
        f"Excavation Method: {proposal.excavation_method}\n"
        f"Utility Type: {proposal.utility_type}\n"
        f"Traffic Diversion: {proposal.expected_traffic_diversion}\n"
        f"Risk Level: {proposal.risk_level}\n"
        f"Status: {proposal.status.upper()}\n"
    )
    if analysis:
        text += (
            f"AI COMPLIANCE ANALYSIS:\n"
            f"Predicted Risk Level: {analysis.risk_predicted.upper()}\n"
            f"AI Recommendation: {analysis.recommendation.upper()}\n"
            f"AI Explanation: {analysis.explanation}\n"
            f"AI Public Impact Score: {analysis.public_impact_score}/100\n"
        )
        if analysis.compliance_report and "violations" in analysis.compliance_report:
            violations = analysis.compliance_report["violations"]
            if violations:
                text += f"AI Violations Found: {', '.join(violations)}\n"
    
    text += f"ADMIN DECISION HISTORY:\n"
    text += f"Current Status: {proposal.status.upper()}\n"
    text += f"{remarks_str}\n"
    return text

async def index_proposal_in_rag(proposal_id: uuid.UUID, remarks: Optional[str] = None, db_session = None) -> None:
    """
    Indexes or updates a proposal's details, AI analysis report, and decisions in ChromaDB.
    """
    from app.models.proposal import Proposal
    from app.models.ai import AIAnalysis
    from sqlalchemy.future import select
    
    proposal = None
    analysis = None
    
    if db_session is None:
        from app.core.database import AsyncSessionLocal
        async with AsyncSessionLocal() as session:
            result = await session.execute(select(Proposal).where(Proposal.id == proposal_id))
            proposal = result.scalars().first()
            if not proposal:
                return
            result_ai = await session.execute(select(AIAnalysis).where(AIAnalysis.proposal_id == proposal_id))
            analysis = result_ai.scalars().first()
    else:
        result = await db_session.execute(select(Proposal).where(Proposal.id == proposal_id))
        proposal = result.scalars().first()
        if not proposal:
            return
        result_ai = await db_session.execute(select(AIAnalysis).where(AIAnalysis.proposal_id == proposal_id))
        analysis = result_ai.scalars().first()

    text = format_proposal_for_rag(proposal, analysis, remarks)
    doc_id = f"proposal_{proposal.id}"
    
    try:
        collection.delete(ids=[doc_id])
    except Exception:
        pass
        
    collection.add(
        documents=[text],
        metadatas=[{
            "policy_id": str(proposal.id),
            "file_name": f"Proposal: {proposal.purpose}",
            "extraction_method": "text",
            "page_number": 1,
            "type": "proposal_decision"
        }],
        ids=[doc_id]
    )
    print(f"RAG: Successfully indexed proposal {proposal.id} details and decisions in ChromaDB.")

async def query_proposal_rag_knowledge_base(proposal_id: uuid.UUID, question: str) -> RAGQueryResponse:
    """
    Queries ChromaDB utilizing both the specific proposal context and global policy circulars.
    """
    # 1. Fetch proposal document
    proposal_doc = ""
    doc_id = f"proposal_{proposal_id}"
    try:
        p_res = collection.get(ids=[doc_id], include=["documents", "metadatas"])
        if p_res and p_res.get("documents"):
            proposal_doc = p_res["documents"][0]
    except Exception as e:
        print(f"Error fetching proposal doc for RAG query: {e}")

    # 2. Query general RAG policy circulars
    query_results = collection.query(
        query_texts=[question],
        n_results=5
    )
    documents = query_results.get("documents", [[]])[0]
    metadatas = query_results.get("metadatas", [[]])[0]

    # 3. Synthesize context
    context_blocks = []
    if proposal_doc:
        context_blocks.append(f"SPECIFIC PLAN DETAILS:\n{proposal_doc}\n")
        
    sources = []
    for doc, meta in zip(documents, metadatas):
        # skip proposal chunk if it matches standard results
        if meta.get("type") == "proposal_decision" or str(proposal_id) in meta.get("policy_id", ""):
            continue
        raw_doc_name = meta.get("file_name", "Unknown Policy File")
        doc_name = clean_policy_filename(raw_doc_name)
        page_num = meta.get("page", 1)
        context_blocks.append(f"Source Circular: {doc_name} (Page {page_num})\nContent:\n{doc}\n")
        sources.append(SourceCitation(
            document_name=doc_name,
            page=page_num,
            excerpt=doc[:150] + "...",
            extraction_method=meta.get("extraction_method", "text")
        ))

    context = "\n".join(context_blocks)

    # 4. Formulate Prompt
    system_instruction = (
        "You are the GCC excavation compliance assistant. You help planners understand compliance reports and admin decisions.\n"
        "Use the SPECIFIC PLAN DETAILS and matching Source Circulars provided below to answer the user's question.\n"
        "If the user asks why the admin rejected the plan, refer specifically to the Status and Admin Decision Remarks in the PLAN DETAILS.\n"
        "Be direct, polite, and cite constraints or policies from the circulars if they are relevant."
    )
    
    user_prompt = (
        f"CONTEXT EXCERPTS:\n{context}\n\n"
        f"USER QUESTION: {question}\n\n"
        f"Provide a natural language answer citing relevant information."
    )

    answer = "I could not retrieve enough details to answer your question."
    try:
        answer = await query_llm(prompt=user_prompt, system_prompt=system_instruction)
    except Exception as e:
        print(f"LLM Error: {e}")
        if "remarks" in context.lower():
            for line in context.split("\n"):
                if "admin decision remarks" in line.lower():
                    answer = f"The admin made the following decision: {line}. Please revise accordingly."
                    break

    return RAGQueryResponse(
        answer=answer,
        sources=sources
    )

def clean_policy_filename(name: str) -> str:
    # If name has a UUID prefix followed by '_', strip it (UUID is 36 chars)
    if len(name) > 37 and name[8] == "-" and name[13] == "-" and name[18] == "-" and name[23] == "-" and name[36] == "_":
        return name[37:]
    return name

async def query_rag_knowledge_base(question: str) -> RAGQueryResponse:
    """
    Runs similarity query in ChromaDB and forwards matching contexts to the failover LLM to synthesize answer.
    """
    # 1. Similarity query top 10 matches
    query_results = collection.query(
        query_texts=[question],
        n_results=10
    )

    documents = query_results.get("documents", [[]])[0]
    metadatas = query_results.get("metadatas", [[]])[0]

    if not documents:
        return RAGQueryResponse(
            answer="No relevant government policy documents found in the GCC knowledge base. Please check back later.",
            sources=[]
        )

    # 2. Format Context Excerpts
    context_blocks = []
    sources = []
    for doc, meta in zip(documents, metadatas):
        raw_doc_name = meta.get("file_name", "Unknown Policy File")
        doc_name = clean_policy_filename(raw_doc_name)
        page_num = meta.get("page", 1)
        
        context_blocks.append(f"Source: {doc_name} (Page {page_num})\nContent:\n{doc}\n")
        sources.append(SourceCitation(
            document_name=doc_name,
            page=page_num,
            excerpt=doc[:200] + "..." if len(doc) > 200 else doc
        ))

    context_str = "\n---\n".join(context_blocks)

    # 3. Formulate RAG strict prompt
    system_instruction = (
        "You are the Greater Chennai Corporation (GCC) Official AI Chat Assistant.\n"
        "Your task is to answer user queries regarding road-cut regulations, guidelines, circulars, and SOPs.\n"
        "RULES:\n"
        "1. Answer ONLY using the provided Policy Context excerpts.\n"
        "2. Do not make up facts or deviate outside the scope of the provided policy documents (no hallucinations).\n"
        "3. You must explicitly cite the document name and page number for the facts you mention.\n"
        "4. If the provided context does not contain enough information to answer the question, state politely "
        "that the information is not present in the indexed GCC policy documents.\n"
        "5. Structure your response professionally: use bold headers, clean bullet points, or numbered lists "
        "for readability. Maintain a clear, administrative tone."
    )

    prompt = (
        f"POLICY CONTEXT:\n"
        f"{context_str}\n\n"
        f"USER QUESTION:\n"
        f"{question}\n\n"
        f"GCC Chat Assistant Answer:"
    )

    # 4. Invoke Failover LLM
    answer = await query_llm(prompt=prompt, system_prompt=system_instruction)

    return RAGQueryResponse(
        answer=answer,
        sources=sources
    )

# ----------------------------------------------------
# HISTORICAL SIMILAR PROJECTS RETRIEVAL
# ----------------------------------------------------
history_collection = chroma_client.get_or_create_collection(
    name="gcc_projects_history",
    embedding_function=embedding_fn
)

def seed_historical_projects():
    """Seeds realistic historical road excavation projects from Chennai to test similarity retrieval."""
    try:
        count = len(history_collection.get().get("ids", []))
        if count > 0:
            return
            
        print("RAG: Seeding historical road-cut project registry...")
        historical_docs = [
            # Project 1
            "Deep trenching and utility relocation for corridor construction of Metro Rail Phase 2 near Anna Salai.",
            # Project 2
            "Excavation for laying 300mm water supply pipes along the IT corridor (OMR) near Velachery junction.",
            # Project 3
            "Underground cable laying for 110KV substation connectivity on Poonamallee High Road."
        ]
        historical_metadatas = [
            {
                "road_name": "Anna Salai",
                "purpose": "Metro Rail Construction & Utility Relocation",
                "outcome": "Successfully completed. Minor traffic bottlenecks resolved using off-peak shifts.",
                "lessons_learned": "Night shift work (11 PM to 5 AM) is critical to prevent traffic gridlocks on arterial roads. Multi-department coordination saved 15% road restoration costs."
            },
            {
                "road_name": "Rajiv Gandhi Salai (OMR)",
                "purpose": "Water Board Pipe Laying",
                "outcome": "Delayed due to overlapping telecom cable cuts.",
                "lessons_learned": "Ground Penetrating Radar (GPR) scanning must be done before trenching to avoid hitting unmapped telecom cabling. Joint coordination could have prevented cable damage."
            },
            {
                "road_name": "Poonamallee High Road",
                "purpose": "Electricity Board Substation Cabling",
                "outcome": "Completed. Heavy rainfall caused minor erosion in trenches.",
                "lessons_learned": "Backfilling and quick road curing must happen within 24 hours of trenching. Avoid work in monsoon season (Oct-Dec) to prevent water logging in open trenches."
            }
        ]
        historical_ids = ["hist_proj_1", "hist_proj_2", "hist_proj_3"]
        
        history_collection.add(
            documents=historical_docs,
            metadatas=historical_metadatas,
            ids=historical_ids
        )
        print("RAG: Successfully seeded Chennai project history collection in ChromaDB.")
    except Exception as e:
        print(f"RAG Warning: Seeding project history failed: {e}")

# Run seeding on startup
seed_historical_projects()

@traceable(run_type="retriever", name="ChromaDB Project Retrieval")
async def retrieve_similar_projects(description: str) -> List[Dict[str, Any]]:
    """
    Retrieves previous excavation projects from the vector database using similarity searches.
    """
    try:
        results = history_collection.query(
            query_texts=[description],
            n_results=2
        )
        
        documents = results.get("documents", [[]])[0]
        metadatas = results.get("metadatas", [[]])[0]
        distances = results.get("distances", [[]])[0]
        
        similar_projects = []
        for i in range(len(documents)):
            # Convert L2 distance to a standard similarity percentage score
            dist = distances[i] if i < len(distances) else 0.5
            similarity_score = max(min(round((1.0 - dist) * 100, 2), 100.0), 0.0)
            
            meta = metadatas[i]
            similar_projects.append({
                "similarity_score": similarity_score,
                "road_name": meta.get("road_name"),
                "purpose": meta.get("purpose"),
                "outcome": meta.get("outcome"),
                "lessons_learned": meta.get("lessons_learned"),
                "description_excerpt": documents[i]
            })
            
        return similar_projects
    except Exception as e:
        print(f"Error retrieving similar projects: {e}")
        return []

